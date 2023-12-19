from docx import Document
from datetime import datetime
from data_collection_util import collect_data_master, collect_data
from generate_time import generate_time
from docx.shared import Pt
import argparse

 # dd-mm-yyyy format
current_date = datetime.now().strftime("%d-%m-%Y")
# day
current_day = datetime.now().strftime("%A")

def _print_table_contents(table):
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            print(f"Row {row_idx}, Column {col_idx}: {cell.text}")

def _update_table_cell(table, row_index, col_index, new_value):
    """
        Row 1, Column 1: DATE
        Row 1, Column 3: DAY
        Row 2, Column 0: SECURITY OFFICER
        Row 2, Column 2: LICENSE NO.
        Row 3, Column 0: START TIME
        Row 3, Column 2: FINISH TIME

    """
    table.cell(row_index, col_index).text = new_value

# collect data based on terminal flag
def get_user_data(args):
    if args.master:
        print("============= Master Function Enabled =============")
        print("- Using predefined officer name and license no.")
        print("============= Welcom Ad =============")
        return collect_data_master()
    else:
        return collect_data()


def update_template(template_path):

    # parser
    parser = argparse.ArgumentParser(description="Use the master function")
    parser.add_argument("-m", "--master", action="store_true", help="Use collect_data_master()")
    args = parser.parse_args()

    # collect data
    data = get_user_data(args)

    # template
    template_doc = Document(template_path)

    # Update date and day in the first table
    first_table = template_doc.tables[0]
    # _print_table_contents(first_table)
    _update_table_cell(first_table, 1, 1, current_date)
    _update_table_cell(first_table, 1, 3, current_day)
    _update_table_cell(first_table, 2, 1, data["officer_name"])
    _update_table_cell(first_table, 2, 3, data["license_no"])

    # second table 
    second_table = template_doc.tables[1]
    # _print_table_contents(second_table)

    # delete the rows except the first one
    for row in second_table.rows[1:]:
        second_table._element.remove(row._element)

    # generate rows
    times = generate_time(data["start_time"], data["finish_time"], data["patrol_times"])
    print(times)
    is_first_row = True
    for time in times:
        # new row
        new_row = second_table.add_row()

        # set the value in the first column
        run_1 = new_row.cells[0].paragraphs[0].add_run()
        font_size = Pt(13)
        run_1.font.size = font_size
        run_1.text = time

        # set the value in the second column
        run_2 = new_row.cells[1].paragraphs[0].add_run()
        run_2.font.size = font_size

        if is_first_row:
            run_2.text = f"{data['officer_name']} on-site, {data['replaced_officer']} left premises"
        else:
            if int(time) % 100 == 0 or int(time) % 100 == 30:
                run_2.text = "Completed patrolling, checked all external doors. All doors are locked, nothing to report"
            else:
                run_2.text = f"MBN Patrol car completed external patrol at {time}-{str(int(time)+2).zfill(4)}"

        is_first_row = False

    # Save the modified document
    template_doc.save(f"{data['serial_no']}Sum Report {current_date} ({str(data['start_time']).zfill(4)}-{str(data['finish_time']).zfill(4)}).docx")
    template_doc.save(f"{current_date}.docx")

if __name__ == "__main__":
    # template document
    template_path = "template.docx"

    # Update the template and save the modified document
    update_template(template_path)

    print("============= Report Generated =============")
